"""
Generate AGIS_AI_AUTHORING_PROPOSAL.docx from the markdown source.
Run: python make_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = 'Arial'
BODY_SIZE = Pt(11)
CODE_SIZE = Pt(9.5)

MD_FILE = "AGIS_AI_AUTHORING_PROPOSAL.md"
OUT_FILE = "AGIS_AI_AUTHORING_PROPOSAL.docx"

# ─── helpers ────────────────────────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Add a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_para(doc, raw_text, style='Normal', level=0):
    """
    Add a paragraph that may contain **bold** and `code` spans.
    level: indent level for bullets (0 = no indent, 1 = one level, etc.)
    """
    p = doc.add_paragraph(style=style)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)

    # Split text on **bold** and `code` markers
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', raw_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.bold = True
        else:
            p.add_run(part)
    return p


# ─── parser ─────────────────────────────────────────────────────────────────

def set_style_font(style, font_name):
    """Force a specific font on a style, overriding theme fonts."""
    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_doc_font(doc, font_name, size):
    """Set Arial on every style that might otherwise show Calibri."""
    target_styles = [
        'Normal', 'Default Paragraph Font',
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'Title', 'Subtitle',
        'List Bullet', 'List Number',
    ]
    for name in target_styles:
        try:
            s = doc.styles[name]
            set_style_font(s, font_name)
            if name == 'Normal':
                s.font.size = size
        except KeyError:
            pass


def generate_docx(md_path, out_path):
    doc = Document()
    set_doc_font(doc, BODY_FONT, BODY_SIZE)

    # Set page margins
    for section in doc.sections:
        section.left_margin  = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin   = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    bullet_buffer = []   # collect consecutive bullet lines

    def flush_bullets():
        for bl in bullet_buffer:
            add_inline_para(doc, bl, style='List Bullet', level=1)
        bullet_buffer.clear()

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip('\n')
        stripped = raw.strip()

        # ── code block fences ─────────────────────────────────────────────
        if stripped.startswith('```'):
            if not in_code_block:
                flush_bullets()
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # Emit code lines indented, same font as body
                for cl in code_lines:
                    p = doc.add_paragraph(cl)
                    p.paragraph_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.font.name = BODY_FONT
                        run.font.size = CODE_SIZE
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── headings ──────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            flush_bullets()
            depth = len(m.group(1))
            text  = m.group(2)
            # Map ### → Heading 3, ## → Heading 2, # → Heading 1
            heading_level = min(depth, 4)
            p = doc.add_heading(level=heading_level)
            # Clear default runs and re-add with inline parsing
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # ── horizontal rule ───────────────────────────────────────────────
        if stripped in ('---', '***', '___'):
            flush_bullets()
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── blockquote ────────────────────────────────────────────────────
        if stripped.startswith('>'):
            flush_bullets()
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            # Light grey border on left
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '6')
            left.set(qn('w:color'), '888888')
            pBdr.append(left)
            pPr.append(pBdr)
            # Inline formatting
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = p.add_run(part)
                    run.italic = True
            i += 1
            continue

        # ── numbered list ─────────────────────────────────────────────────
        m_num = re.match(r'^\d+\.\s+(.*)', stripped)
        if m_num:
            flush_bullets()
            text = m_num.group(1)
            add_inline_para(doc, text, style='List Number', level=1)
            i += 1
            continue

        # ── bullet list ───────────────────────────────────────────────────
        m_bullet = re.match(r'^[-*+]\s+(.*)', stripped)
        if m_bullet:
            text = m_bullet.group(1)
            bullet_buffer.append(text)
            i += 1
            continue

        # ── empty line ────────────────────────────────────────────────────
        if stripped == '':
            flush_bullets()
            # Only add a spacer paragraph if the next line isn't also blank
            next_stripped = lines[i+1].strip() if i+1 < len(lines) else ''
            if next_stripped != '' and not next_stripped.startswith('#'):
                doc.add_paragraph()
            i += 1
            continue

        # ── normal paragraph ──────────────────────────────────────────────
        flush_bullets()
        add_inline_para(doc, stripped, style='Normal')
        i += 1

    flush_bullets()

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    generate_docx(MD_FILE, OUT_FILE)
